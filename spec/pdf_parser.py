"""
Document text extraction from PDF and Word files.
"""

from pathlib import Path
import PyPDF2
import pdfplumber


def extract_full_text(file_path: str) -> str:
    """Extract complete text from a PDF or Word document."""
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document file not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == '.pdf':
        return _extract_pdf_text(file_path)
    elif suffix in ['.docx', '.doc']:
        return _extract_word_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: .pdf, .docx, .doc")


def _extract_pdf_text(pdf_path: str) -> str:
    """Extract text from PDF file."""
    text_parts = []

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception:
        with open(pdf_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

    full_text = '\n\n'.join(text_parts)

    if not full_text.strip():
        raise ValueError(f"No text extracted from PDF: {pdf_path}")

    return full_text


def _extract_word_text(word_path: str) -> str:
    """Extract text from Word document (.docx or .doc)."""
    from docx import Document

    doc = Document(word_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    full_text = '\n\n'.join(text_parts)

    if not full_text.strip():
        raise ValueError(f"No text extracted from Word document: {word_path}")

    return full_text
